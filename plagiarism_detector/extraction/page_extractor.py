# -*- coding: utf-8 -*-
"""
وحدة استخراج نصوص المستندات مع الحفاظ التام على أرقام الصفحات (Page-Level Document Extractor):
- استخراج صفحة-بصفحة من ملفات PDF مع حفظ رقم الصفحة الأصلي بدقة.
- استخراج فقرات مستندات Word (.docx) و (.txt) مع وسم رقم الصفحة بـ None (غير متاح).
- دعم معالجة الحروف المعكوسة (Bidi) وتنظيف شوائب خطوط PDF العربية.
- تفعيل التعرف الضوئي (OCR) المشروط فقط على الصفحات المصمتة التي يقل نصها عن 20 كلمة.
"""

import os
import re
import logging
import unicodedata
from pathlib import Path

from .ocr_engine import check_ocr_availability, ocr_page_pixmap

logger = logging.getLogger(__name__)

# الكلمات المعكوسة الشائعة لاكتشاف مشاكل الـ Bidi
_REVERSED_ARABIC = re.compile(r'\b(ة[ا-ي]{2,}|[ا-ي]{2,}لا|[ا-ي]{2,}ةي|هيلع|ءاكذ|ميلعت)\b')
_ARABIC_CHAR_RE = re.compile(r'[\u0600-\u06FF\uFB50-\uFDFF\uFE70-\uFEFF]')
_LATIN_CHAR_RE = re.compile(r'[a-zA-Z]')


def clean_arabic_pdf_glyphs(text: str) -> str:
    """
    تنظيف الشوائب الناتجة عن فك ترميز خطوط PDF العربية:
    - إزالة المحارف الصفرية المعطوبة.
    - تصحيح الشوائب الشهيرة مثل (امل / املم -> ال).
    - تصحيح بعض الكلمات المعكوسة شائعة الحدوث في عناوين الخطوط القديمة.
    """
    if not text:
        return ''

    text = unicodedata.normalize('NFKC', text)
    # إزالة الرموز والمحارف التحكمية غير المقروءة
    text = re.sub(r'[\uFFFD\u0000-\u0008\u000B\u000C\u000E-\u001F\u25C6\u25C7\u25A0\u25A1]', ' ', text)
    text = text.replace('\u0640', '')  # إزالة الكشيدة

    # تصحيح تراكيب ال التعريف الشائعة في الخطوط المعطوبة
    text = re.sub(r'\bاملم([ا-ي]{2,})', r'الم\1', text)
    text = re.sub(r'\bامل([ا-ي]{2,})', r'الم\1', text)
    text = re.sub(r'\bيف\b', 'في', text)
    text = re.sub(r'\bعل[»|م]\b', 'على', text)
    text = re.sub(r'\bإل[»|م]\b', 'إلى', text)

    return text


def unreverse_arabic_text(text: str) -> str:
    """
    كشف وتصحيح الكلمات العربية المعكوسة من ملفات الـ PDF القديمة مع تجنب المساس بالنصوص الإنجليزية والفرنسية.
    """
    if not text:
        return ''

    lines = text.split('\n')
    out_lines = []

    for line in lines:
        if _LATIN_CHAR_RE.search(line):
            out_lines.append(line)
            continue

        words = line.split()
        if not words:
            out_lines.append(line)
            continue

        arabic_words = [w for w in words if _ARABIC_CHAR_RE.search(w)]
        if len(arabic_words) < 3:
            out_lines.append(line)
            continue

        rev_count = sum(1 for w in arabic_words if w.startswith('ة') or w.endswith('لا') or w.startswith('ات'))
        if (rev_count / len(arabic_words)) >= 0.55:
            # ترتيب الكلمات والحروف معكوس
            fixed = []
            for w in reversed(words):
                fixed.append(w[::-1])
            out_lines.append(' '.join(fixed))
        else:
            out_lines.append(line)

    return '\n'.join(out_lines)


def extract_document_pages(file_path: str, enable_ocr: bool = True) -> list[dict]:
    """
    استخراج محتوى المستند صفحة-بصفحة.
    يرجع قائمة بالقواميس:
    [
        {'page_number': 1, 'text': '...', 'is_ocr': False},
        ...
    ]
    في ملفات Word و TXT، تكون قيمة 'page_number': None
    """
    ext = os.path.splitext(file_path)[1].lower()
    pages: list[dict] = []

    if ext == '.pdf':
        pages = _extract_pages_from_pdf(file_path, enable_ocr=enable_ocr)
    elif ext in ('.docx', '.doc'):
        pages = _extract_pages_from_docx(file_path)
    elif ext == '.txt':
        pages = _extract_pages_from_txt(file_path)
    else:
        logger.warning(f"صيغة ملف غير مدعومة للاستخراج: {ext}")

    return pages


def _extract_pages_from_pdf(path: str, enable_ocr: bool = True) -> list[dict]:
    """استخراج صفحات PDF مع الحفاظ على ترقيم الصفحات ومحاولة OCR المشروط."""
    pages = []
    ocr_status = check_ocr_availability() if enable_ocr else {'available': False}

    # المحاولة الأولى: PyMuPDF (fitz)
    try:
        import fitz
        doc = fitz.open(path)
        for idx, page in enumerate(doc):
            page_num = idx + 1
            text = page.get_text('text')
            is_ocr = False

            # فحص إذا كانت الصفحة مصمتة أو مسحوبة عبر سكانر (أقل من 20 كلمة)
            words = text.split() if text else []
            if len(words) < 20 and enable_ocr and ocr_status['available']:
                logger.info(f"الصفحة {page_num} نصها قليل جداً ({len(words)} كلمة)، جاري تجربة OCR مشروط...")
                try:
                    pix = page.get_pixmap(dpi=200)
                    ocr_text = ocr_page_pixmap(pix)
                    if ocr_text and len(ocr_text.split()) > len(words):
                        text = ocr_text
                        is_ocr = True
                except Exception as oe:
                    logger.debug(f"فشل OCR للصفحة {page_num}: {oe}")

            if text:
                text = unicodedata.normalize('NFKC', text)
                text = unreverse_arabic_text(text)
                text = clean_arabic_pdf_glyphs(text)

            pages.append({
                'page_number': page_num,
                'text': text.strip() if text else '',
                'is_ocr': is_ocr
            })
        return pages
    except Exception as e:
        logger.warning(f"تعذر استخراج PDF بـ PyMuPDF: {e}. جاري استخدام pdfplumber...")

    # المحاولة البديلة: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for idx, p in enumerate(pdf.pages):
                page_num = idx + 1
                t = p.extract_text() or ''
                if t:
                    t = unicodedata.normalize('NFKC', t)
                    t = unreverse_arabic_text(t)
                    t = clean_arabic_pdf_glyphs(t)
                pages.append({
                    'page_number': page_num,
                    'text': t.strip(),
                    'is_ocr': False
                })
        return pages
    except Exception as e2:
        logger.error(f"فشل استخراج ملف PDF بالكامل: {e2}")
        return []


def _extract_pages_from_docx(path: str) -> list[dict]:
    """استخراج نصوص Word مع التوثيق الصريح بأن رقم الصفحة غير متاح."""
    try:
        import docx
        doc = docx.Document(path)
        full_paragraphs = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                full_paragraphs.append(clean_arabic_pdf_glyphs(p.text.strip()))

        # استخراج نصوص الجداول أيضاً
        for tbl in doc.tables:
            for row in tbl.rows:
                row_txt = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_txt:
                    full_paragraphs.append(' | '.join(row_txt))

        joined = '\n\n'.join(full_paragraphs)
        return [{
            'page_number': None,  # غير متاح لملفات Word
            'text': joined,
            'is_ocr': False
        }]
    except Exception as e:
        logger.error(f"فشل استخراج ملف DOCX: {e}")
        return []


def _extract_pages_from_txt(path: str) -> list[dict]:
    """استخراج النصوص العادية."""
    try:
        for enc in ['utf-8', 'utf-8-sig', 'windows-1256', 'latin-1']:
            try:
                with open(path, 'r', encoding=enc) as f:
                    text = f.read()
                return [{
                    'page_number': None,
                    'text': clean_arabic_pdf_glyphs(text),
                    'is_ocr': False
                }]
            except UnicodeDecodeError:
                continue
    except Exception as e:
        logger.error(f"فشل استخراج ملف TXT: {e}")
    return []


def extract_text(file_path: str) -> str:
    """
    دالة للتوافق العكسي الكامل مع المستدعين القدامى: ترجع النص الكامل للمستند كسلسلة نصية واحدة.
    """
    pages = extract_document_pages(file_path, enable_ocr=False)
    return '\n\n'.join(p['text'] for p in pages if p['text'])
